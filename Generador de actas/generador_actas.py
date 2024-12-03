import tkinter as tk
from tkinter import messagebox
from tkinter import ttk  # Para usar Combobox
import pandas as pd
from openpyxl import load_workbook, Workbook
from openpyxl.worksheet.protection import SheetProtection
from openpyxl.workbook.protection import WorkbookProtection

import sys
import os

# Función para obtener la ruta del recurso
def obtener_ruta_recurso(relativa):
    if getattr(sys, 'frozen', False):  # Si está empaquetado con PyInstaller
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relativa)

# Ejemplo de uso
ruta_permiso = obtener_ruta_recurso("recursos\PERMISOS EXAMEN - 2023.docx")
ruta_acta = obtener_ruta_recurso("recursos\ACTA DE EXAMEN.docx")

# Crear el DataFrame inicial vacío
columnas = ["N°","ALUMNO", "DNI", "MODALIDAD", "CONDICION", "CURSO", "ESPACIO CURRICULAR"]

# Intentar leer el archivo Excel
archivo_excel = "acta_de_examen.xlsx"
contraseña = "153570"
try:
    df = pd.read_excel(archivo_excel, dtype={"DNI": int})  # Leer el archivo Excel
except FileNotFoundError:
    # Si el archivo no existe, crear uno nuevo con las columnas predeterminadas
    df = pd.DataFrame(columns=columnas)
    with pd.ExcelWriter(archivo_excel, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="PERMISOS", index=False)  # Crear el archivo Excel con las columnas iniciales
        df.to_excel(writer, sheet_name="MMO", index=False)
        df.to_excel(writer, sheet_name="TEM", index=False)
    # Proteger el archivo Excel
    wb = load_workbook(archivo_excel)
    ws = wb.active
    
    # Proteger la hoja
    ws.protection = SheetProtection(sheet=True, password=contraseña)
    
    # Proteger el libro completo
    wb.security = WorkbookProtection(workbookPassword=contraseña, lockStructure=True)
    wb.save(archivo_excel)

# Listas de materias para las modalidades MMO y TEM
materias_mmo = [
    "LENGUA Y LITERATURA I", "LENGUA EXTRANJERA I", "HISTORIA I", "GEOGRAFÍA I",
    "FORMACIÓN ETICA Y CIUDADANA I", "EDUC. ARTÍSTICA: MÚSICA", "EDUCACIÓN FÍSICA I",
    "MATEMATICA I", "BIOLOGÍA I", "FÍSICA QUÍMICA I", "TECNOLOGÍA DE LOS MATERIALES",
    "DIBUJO TÉCNICO I", "TALLER I: CARP. ELECT. HERRERIA", "LENGUA Y LITERATURA II",
    "LENGUA EXTRANJERA II", "HISTORIA II", "GEOGRAFÍA II", "FORMACIÓN ETICA Y CIUDADANA II",
    "EDUC. ARTÍSTICA: ARTES VISUALES", "EDUCACIÓN FÍSICA II", "MATEMATICA II", "BIOLOGÍA II",
    "FÍSICA II", "QUIMICA II", "TECNOLOGÍA DE LOS PROCESOS", "DIBUJO TÉCNICO II", "TALLER II",
    "LENGUA Y LITERATURA III", "LENGUA EXTRANJERA III", "HISTORIA III", "GEOGRAFÍA III",
    "FORMACIÓN ETICA Y CIUDADANA III", "EDUC. ARTÍSTICA: TEATRO/ DANZA", "EDUCACIÓN FÍSICA III",
    "MATEMATICA III", "BIOLOGÍA III", "FÍSICA III", "QUIMICA III", "TIC`S", "REPRES.GRAFICA - CAD",
    "TALLER III - CONSTRUCCIONES", "LENGUA Y LITERATURA IV", "LENGUA EXTRANJERA IV",
    "HISTORIA DE CATAMARCA Y EL NOA", "GEOGRAFÍA DE CATAMARCA Y EL NOA", "EDUCACIÓN FÍSICA",
    "MATEMATICA IV", "FÍSICA IV", "QUÍMICA IV", "GEOMETRÍA DESCRIPTIVA",
    "ESTATICA Y RESISTENCIA DE LOS MATERIALES", "SISTEMAS CONSTRUCTIVOS I", "TALLER IV- CONSTRUCCIONES",
    "LENGUA Y LITERATURA V", "INGLÉS TÉCNICO I", "CONSTRUCCION CIUDADANA - ESI", "EDUCACIÓN FÍSICA",
    "MATEMATICA V", "QUÍMICA APLICADA", "SISTEMAS CONSTRUCTIVO II", "ARQUITECTURA", "DISEÑO Y CALCULO DE ESTRUCTURAS",
    "PROYECTO I", "TALLER V - TERMINACIONES HºAº", "LENGUA Y LITERATURA VI", "INGLÉS TÉCNICO II",
    "EDUCACIÓN FÍSICA", "MATEMATICA VI", "MARCO JURÍDICO DE LOS PROC. CONSTRUCTIVOS",
    "CÓMPUTOS Y PRESUPUESTOS", "TOPOGRAFÍA", "SISTEMAS ELECTRICOS Y ELECTRÓNICOS DOMICILIARIOS",
    "SISTEMAS DE MECANIZADO CNC", "SISMORRESISTENTE CONSTRUCCIONES", "GESTION DE OBRAS", "PROYECTO FINAL"
]

materias_tem = [
    "FORMACIÓN ETICA Y CIUDADANA III", "EDUC. ARTÍSTICA: TEATRO/ DANZA", "EDUCACIÓN FÍSICA III",
    "MATEMATICA III", "BIOLOGÍA III", "FÍSICA III", "QUIMICA III", "TIC`S", "DIBUJO III",
    "TALLER III - ELECT- MECANICA - HERRERIA", "LENGUA Y LITERATURA IV", "LENGUA EXTRANJERA IV",
    "EDUCACIÓN FÍSICA", "GEOGRAFÍA DE CATAMARCA Y EL NOA", "HISTORIA DE CATAMARCA Y EL NOA",
    "ANALISIS MATEMATICO I", "FÍSICA IV", "QUÍMICA APLICADA", "TRANSF. Y SINTESIS DE LOS MATERIALES",
    "MAQUINAS ELÉCTRICAS Y AUTOMATISMOS I", "DISEÑO Y PROCESAM MECÁNICO I", "INST. Y APLIC DE LA ENERGÍA I",
    "ELECTRICIDAD", "MECÁNICA", "METALURGIA", "LENGUA Y LITERATURA", "INGLES TECNICO I", "EDUCACION FÍSICA",
    "CONSTRUCCION CIUDADANA - ESI", "ANÁLISIS MATEMATICO II", "MECÁNICA Y MECANISMOS",
    "RESIST. Y ENSAYO DE LOS MATERIALES", "LAB. DE MEDICIONES ELÉCTRICAS", "MÁQUINAS ELÉCTRICAS Y AUTOMATISMOS II",
    "DISEÑO Y PROCESAM MECÁNICO II", "ELECTROTECNIA I", "INST. Y APLIC DE LA ENERGÍA II", "REDES ELÉCTRICAS",
    "MECÁNICA", "SOLDADURA", "LENGUA Y LITERATURA", "INGLÉS TECNICO II", "EDUCACIÓN FÍSICA",
    "MATEMATICA APLICADA", "TERMOD. Y MÁQ. TÉRMICAS", "GESTIÓN Y ADMIN. INDUSTRIAL",
    "MÁQUINAS ELÉCTRICAS Y AUTOMATISMOS III", "DISEÑO Y PROCESAM. MECÁNICO III", "INST. Y APLICAC. DE LA ENERGÍA III",
    "SISTEMAS MECÁNICOS I", "CONTROL Y AUTOMATISMOS INDUSTRIALES", "ELECTROTECNIA II", "SOLDADURA II",
    "INST. DE CONTROL AUTOMATIZADO", "SISTEMAS DE MECANIZADO", "COMUNICACIÓN", "EDUCACIÓN FÍSICA",
    "GESTION DE PYMES INDUSTRIALES", "DERECHO DEL TRABAJO", "LAB. DE METROLOGÍA Y CONTROL DE CALIDAD",
    "MONTAJE ELECTROMECÁNICO", "MANTENIMIENTO INDUSTRIAL", "SEGURIDAD, HIGIENE Y PROT. AMBIENTAL",
    "SISTEMAS MECÁNICOS II", "ELECTRÓNICA INDUSTRIAL", "TALLER DE ELECTRÓNICA", "PROY. Y DISEÑO ELECTROMECÁNICO",
    "PROY. Y DISEÑO DE INST. ELÉCTRICAS", "SISTEMAS DE MECANIZADO CNC", "PASANTÍA"
]

materias_mmo_con_sufijo = [f"{materia} (MMO)" for materia in materias_mmo]
materias_tem_con_sufijo = [f"{materia} (TEM)" for materia in materias_tem]


from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx import Document


def aplicar_estilo_encabezado(paragraph, texto, texto1, fuente="Calibri", tamaño=11, negrita=False):
    """
    Aplica el estilo Calibri de tamaño 11 solo a los encabezados, sin alterar el formato.
    """
    # Reemplazamos el texto sin cambiar el formato
    for run in paragraph.runs:
        if texto in run.text:  # Buscar el texto a reemplazar
            run.text = run.text.replace(run.text, texto1)
            run.font.name = fuente
            run.font.size = Pt(tamaño)
            run.font.bold = negrita

def aplicar_estilo_personalizado_celda(celda, texto, fuente="Arial", tamaño=11, negrita=False):
    """
    Escribe texto en un párrafo dentro de una celda con estilo personalizado.
    """
    # Aseguramos que estamos trabajando con un párrafo en la celda
    p = celda.paragraphs[0]  # Primer párrafo en la celda
    p.clear()  # Limpiar el contenido previo
    run = p.add_run(texto)  # Agregar texto al párrafo
    run.font.name = fuente  # Tipo de letra
    run.font.size = Pt(tamaño)  # Tamaño de letra
    run.bold = negrita  # Negrita

    # Configuración explícita para garantizar el tipo de letra
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), fuente)
    rFonts.set(qn('w:hAnsi'), fuente)
    rFonts.set(qn('w:eastAsia'), fuente)
    rFonts.set(qn('w:cs'), fuente)
    rPr.append(rFonts)

import re  # Importar para manejar caracteres especiales

def generar_actas_excel(archivo_excel, modalidad):

    # Cargar el archivo Excel
    wb = load_workbook(archivo_excel)
    ws = wb.active

    # Desbloquear la hoja para ediciones
    ws.protection = SheetProtection(sheet=False)

    # Leer la hoja correspondiente (TEM o MMO) del archivo Excel
    df = pd.read_excel(archivo_excel, sheet_name=modalidad)

    # Obtener las columnas dinámicas de alumnos y DNIs
    columnas_alumno = [col for col in df.columns if col.startswith("ALUMNO")]
    columnas_dni = [col for col in df.columns if col.startswith("DNI")]

    # Agrupar por materia y curso
    grupos = df.groupby(["ESPACIO CURRICULAR", "CURSO", "CONDICION"])

    for (materia, curso, condicion), grupo in grupos:
        # Crear un documento basado en la plantilla
        doc = Document(ruta_acta)

        # Llenar encabezados del acta
        for paragraph in doc.paragraphs:
            if "EXAMEN DE ALUMNO" in paragraph.text:
                aplicar_estilo_encabezado(paragraph, "REGULAR", condicion)
            if "ESPACIO CURRICULAR" in paragraph.text:
                aplicar_estilo_encabezado(paragraph, "HISTORIA III", materia.split(" (")[0])
            if "PLAN DE ESTUDIO" in paragraph.text:
                aplicar_estilo_encabezado(paragraph, "TEM", modalidad)
            if "CURSO" in paragraph.text:
                aplicar_estilo_encabezado(paragraph, "3º1º", curso)

        # Llenar la tabla con los alumnos
        table = doc.tables[0]
        orden = 1
        fila_actual = 2

        for _, fila in grupo.iterrows():
            for col_alumno, col_dni in zip(columnas_alumno, columnas_dni):
                alumno = fila[col_alumno]
                dni = fila[col_dni]

                if pd.notna(alumno):
                    if fila_actual < len(table.rows):
                        celda = table.rows[fila_actual].cells
                    else:
                        celda = table.add_row().cells

                    aplicar_estilo_personalizado_celda(celda[0], f"{orden:02}", fuente="Arial", tamaño=11, negrita=True)
                    aplicar_estilo_personalizado_celda(celda[2], alumno, fuente="Arial", tamaño=9, negrita=True)
                    aplicar_estilo_personalizado_celda(celda[9], str(int(dni)) if pd.notna(dni) else "", fuente="Arial", tamaño=9, negrita=True)

                    for idx in [1, 3, 4, 5, 6, 7, 8]:
                        aplicar_estilo_personalizado_celda(celda[idx], "", fuente="Arial", tamaño=9, negrita=True)

                    orden += 1
                    fila_actual += 1

        # Limpiar caracteres inválidos en el nombre del archivo
        materia_limpia = re.sub(r'[<>:"/\\|?*]', '', materia.split(" (")[0])
        curso_limpio = re.sub(r'[<>:"/\\|?*]', '', curso)
        modalidad_limpia = re.sub(r'[<>:"/\\|?*]', '', modalidad)
        condicion_limpia = re.sub(r'[<>:"/\\|?*]', '', condicion)

        # Crear la carpeta para la modalidad si no existe
        carpeta_modalidad = "ACTAS_"+modalidad.upper()
        if not os.path.exists(carpeta_modalidad):
            os.makedirs(carpeta_modalidad)

        # Generar el nombre del archivo
        nombre_archivo = f"Acta_{modalidad_limpia}_{materia_limpia}_{curso_limpio}_{condicion_limpia}.docx"

        # Generar la ruta completa
        ruta_archivo = os.path.join(carpeta_modalidad, nombre_archivo)

        # Guardar el documento
        doc.save(ruta_archivo)

    # Reproteger la hoja después de editar
    wb = load_workbook(archivo_excel)
    ws = wb.active
    ws.protection = SheetProtection(sheet=True, password=contraseña)

    # Guardar el archivo nuevamente con protección
    wb.save(archivo_excel)

import openpyxl
from openpyxl.styles import Alignment

from datetime import datetime
import locale

def generar_permisos():
    # Cargar el archivo Excel
    wb = load_workbook(archivo_excel)
    ws = wb.active

    # Desbloquear la hoja para ediciones
    ws.protection = SheetProtection(sheet=False)

    # Leer la hoja correspondiente (TEM o MMO) del archivo Excel
    df = pd.read_excel(archivo_excel, sheet_name='PERMISOS')

    # Obtener las columnas dinámicas de alumnos y DNIs
    columnas_espacio = [col for col in df.columns if col.startswith("ESPACIO CURRICULAR")]
    columnas_curso = [col for col in df.columns if col.startswith("CURSO")]

    # Agrupar por alumno
    grupos = df.groupby(["N°", "ALUMNO", "DNI", "MODALIDAD"])
    # Establecer el idioma a español
    locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')  # En sistemas Linux/Mac
    # locale.setlocale(locale.LC_TIME, 'Spanish_Spain.1252')  # En sistemas Windows

    # Obtener la fecha actual en el formato deseado
    fecha_actual = datetime.now().strftime("TINOGASTA, %d de %B DE %Y").capitalize()

    for (nro, alumno, dni, modalidad), grupo in grupos:
        # Crear un documento basado en la plantilla
        doc = Document(ruta_permiso)

        # Llenar encabezados del acta
        for paragraph in doc.paragraphs:

            if "PERMISO DE EXAMEN N°" in paragraph.text:
                aplicar_estilo_encabezado(paragraph, "56", str(nro), fuente="Arial", tamaño=12, negrita=True)
            if "alumno/a" in paragraph.text:
                aplicar_estilo_encabezado(paragraph, "ORTIZ JOEL ALEXANDER", alumno, fuente="Arial", tamaño=12, negrita=True)
            if "DNI" in paragraph.text:
                aplicar_estilo_encabezado(paragraph, "48663822", str(dni), fuente="Arial", tamaño=12, negrita=True)
            if "Plan de Estudios de" in paragraph.text:
                aplicar_estilo_encabezado(paragraph, "TEM", modalidad, fuente="Arial", tamaño=12, negrita=True)

        # Llenar la tabla con los alumnos
        table = doc.tables[0]
        orden = 1
        fila_actual = 1

        for _, fila in grupo.iterrows():
            for col_espacio, col_curso in zip(columnas_espacio, columnas_curso):
                materia = fila[col_espacio]
                curso = fila[col_curso]

                if pd.notna(materia):
                    if fila_actual < len(table.rows):
                        celda = table.rows[fila_actual].cells
                    else:
                        celda = table.add_row().cells

                    aplicar_estilo_personalizado_celda(celda[0], f"{orden:02}", fuente="Arial", tamaño=11, negrita=True)
                    aplicar_estilo_personalizado_celda(celda[1], materia.split(" (")[0], fuente="Arial", tamaño=11, negrita=True)
                    aplicar_estilo_personalizado_celda(celda[2], str(curso) if pd.notna(dni) else "", fuente="Arial",
                                                       tamaño=11, negrita=True)

                    for idx in [3, 4, 5]:
                        aplicar_estilo_personalizado_celda(celda[idx], "", fuente="Arial", tamaño=11, negrita=True)

                    orden += 1
                    fila_actual += 1
        # Reemplazar el texto específico en los párrafos
        for paragraph in doc.paragraphs:
            if "TINOGASTA, 30 de junio 		DE 2024" in paragraph.text:
                paragraph.text = paragraph.text.replace("TINOGASTA, 30 de junio 		DE 2024", fecha_actual)
                # Modificar el estilo de la fuente
                run = paragraph.runs[0]  # El "run" representa un fragmento del texto en el párrafo
                run.font.name = 'Arial'  # Cambiar la fuente a Arial
                run.font.size = Pt(10)  # Cambiar el tamaño de la fuente
                run.font.bold = True  # Negrita

        # Limpiar caracteres inválidos en el nombre del archivo
        alumno_limpio = re.sub(r'[<>:"/\\|?*]', '', alumno)

        # Crear la carpeta para la modalidad si no existe
        carpeta_alumno = "PERMISO_" + alumno_limpio.upper()
        if not os.path.exists(carpeta_alumno):
            os.makedirs(carpeta_alumno)

        # Generar el nombre del archivo
        nombre_archivo = f"Permiso_{alumno_limpio}.docx"

        # Generar la ruta completa
        ruta_archivo = os.path.join(carpeta_alumno, nombre_archivo)

        # Guardar el documento
        doc.save(ruta_archivo)
    
    # Reproteger la hoja después de editar
    wb = load_workbook(archivo_excel)
    ws = wb.active
    ws.protection = SheetProtection(sheet=True, password=contraseña)

    # Guardar el archivo nuevamente con protección
    wb.save(archivo_excel)

    messagebox.showinfo("Éxito",
                        "Permisos generados correctamente.")


def ajustar_columnas_automatically(archivo_excel, sheet):
    # Cargar el archivo Excel usando openpyxl
    libro = openpyxl.load_workbook(archivo_excel)

    # Seleccionamos la hoja que se pasa como parametro
    hoja = libro[sheet]

    # Recorrer todas las columnas y ajustarlas automáticamente
    for col in hoja.columns:
        max_length = 0
        for cell in col:
            try:
                # Si la celda tiene valor numérico, asegurarse de que se trata como texto
                if isinstance(cell.value, (int, float)):  # Si es un número
                    cell.value = str(cell.value)  # Convertir a cadena (texto)

                # Comprobar la longitud de la cadena (texto) de la celda
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))

                # Opcional: Alinear el texto (para hacer la tabla más legible)
                cell.alignment = Alignment(horizontal='center')

            except:
                pass

        # Ajustar el ancho de la columna, un poco de espacio extra
        adjusted_width = max_length + 2  # Añadir algo de espacio para que no se vea demasiado ajustado
        hoja.column_dimensions[col[0].column_letter].width = adjusted_width

    # Guardar el archivo con los ajustes de tamaño
    libro.save(archivo_excel)


def generar_actas():

    # Cargar el archivo Excel
    wb = load_workbook(archivo_excel)
    ws = wb.active

    # Desbloquear la hoja para ediciones
    ws.protection = SheetProtection(sheet=False)
    try:
        with open(archivo_excel, "r+"):
            pass  # Si podemos abrirlo en modo lectura-escritura, está disponible
    except IOError:
        messagebox.showerror("Error",
                             "El archivo Excel está abierto. Por favor, ciérrelo antes de generar las actas.")
        return

    # Leer archivo de inscripciones
    datos = pd.read_excel(archivo_excel)

    # Reestructurar datos: iterar sobre las columnas dinámicas
    columnas_dinamicas = [col for col in datos.columns if col.startswith("ESPACIO CURRICULAR")]
    datos_reestructurados = []

    for i, columna_espacio in enumerate(columnas_dinamicas, start=1):
        # Extraer columnas relacionadas (Espacio Curricular, Curso, Condición)
        if i == 1:
            columna_curso = "CURSO"
            columna_condicion = "CONDICION"
        else:
            columna_curso = f"CURSO {i}"
            columna_condicion = f"CONDICION {i}"

        # Filtrar filas no vacías en las columnas actuales
        filas_validas = datos[~datos[columna_espacio].isna()]

        # Crear un DataFrame temporal con los valores relevantes
        df_temp = pd.DataFrame({
            "ESPACIO CURRICULAR": filas_validas[columna_espacio],
            "CURSO": filas_validas[columna_curso],
            "CONDICION": filas_validas[columna_condicion],
            "ALUMNO": filas_validas["ALUMNO"],
            "DNI": filas_validas["DNI"]
        })
        datos_reestructurados.append(df_temp)

    # Concatenar todos los DataFrames y filtrar por modalidad
    datos_reestructurados = pd.concat(datos_reestructurados, ignore_index=True)
    datos_mmo = datos_reestructurados[datos_reestructurados["ESPACIO CURRICULAR"].isin(materias_mmo_con_sufijo)].copy()
    datos_tem = datos_reestructurados[datos_reestructurados["ESPACIO CURRICULAR"].isin(materias_tem_con_sufijo)].copy()

    # Añadir columnas de modalidad
    datos_mmo["MODALIDAD"] = "MMO"
    datos_tem["MODALIDAD"] = "TEM"

    # Función para reorganizar los datos según el formato requerido
    def reorganizar_datos(df):
        grupos = df.groupby(["ESPACIO CURRICULAR", "CURSO", "MODALIDAD", "CONDICION"])
        filas = []
        for (espacio, curso, modalidad, condicion), grupo in grupos:
            # Crear una fila para cada grupo
            fila = {
                "ESPACIO CURRICULAR": espacio,
                "CURSO": curso,
                "MODALIDAD": modalidad,
                "CONDICION": condicion
            }
            # Agregar los alumnos y DNIs con índices
            for i, (_, alumno) in enumerate(grupo.iterrows(), start=1):
                fila[f"ALUMNO {i}"] = alumno["ALUMNO"]
                fila[f"DNI {i}"] = alumno["DNI"]
            filas.append(fila)
        return pd.DataFrame(filas)

    # Reorganizar los datos para cada modalidad
    datos_mmo_reorganizados = reorganizar_datos(datos_mmo)
    datos_tem_reorganizados = reorganizar_datos(datos_tem)

    # Guardar en un nuevo archivo Excel
    with pd.ExcelWriter(archivo_excel, engine='openpyxl') as writer:
        datos.to_excel(writer, sheet_name="PERMISOS", index=False)
        datos_mmo_reorganizados.to_excel(writer, sheet_name="MMO", index=False)
        datos_tem_reorganizados.to_excel(writer, sheet_name="TEM", index=False)
    # Reproteger la hoja después de editar
    wb = load_workbook(archivo_excel)
    ws = wb.active
    ws.protection = SheetProtection(sheet=True, password=contraseña)

    # Guardar el archivo nuevamente con protección
    wb.save(archivo_excel)

def generar_actas_imprimir():
    generar_actas()
    hojaMMO = pd.read_excel(archivo_excel, sheet_name="MMO")
    # Verificar si está vacía
    if hojaMMO.empty:
        generar_actas_excel(archivo_excel, "TEM")
    hojaTEM = pd.read_excel(archivo_excel, sheet_name="TEM")
    if hojaTEM.empty:
        generar_actas_excel(archivo_excel, "MMO")
    if hojaMMO.empty == False and hojaTEM.empty == False:
        generar_actas_excel(archivo_excel, "TEM")
        generar_actas_excel(archivo_excel, "MMO")
    messagebox.showinfo("Éxito",
                        "Actas generadas correctamente.")
# Función para registrar los datos
import os
import openpyxl
from openpyxl.utils.exceptions import InvalidFileException

def registrar():
    global df

    # Obtener datos del formulario
    nombre = entrada_nombre.get().strip().upper()
    dni = entrada_dni.get().strip()
    curso = combobox_curso.get().strip().upper()
    especialidad = combobox_especialidad.get().strip().upper()
    materia = f"{combobox_materia.get().strip().upper()} ({especialidad})"
    condicion = combobox_condicion.get().strip().upper()

    # Validar campos obligatorios
    if not nombre or not dni or not curso or not especialidad or not materia or not condicion:
        messagebox.showerror("Error", "Todos los campos son obligatorios.")
        return

    # Validar modalidad y condición
    if especialidad not in ["MMO", "TEM"]:
        messagebox.showerror("Error", "La modalidad debe ser MMO o TEM.")
        return
    if condicion not in ["REGULAR", "LIBRE"]:
        messagebox.showerror("Error", "La condición debe ser LIBRE o REGULAR.")
        return

    dni = str(dni)

    # Verificar si el archivo Excel está abierto
    archivo_excel = "acta_de_examen.xlsx"
    try:
        with open(archivo_excel, "r+"):
            pass  # Si podemos abrirlo en modo lectura-escritura, está disponible
    except IOError:
        messagebox.showerror("Error",
                             "El archivo Excel está abierto. Por favor, ciérrelo antes de registrar al alumno.")
        return
    df = pd.read_excel(archivo_excel)

    # Verificar si el alumno ya existe en el DataFrame
    if int(dni) in df["DNI"].values:
        columnas_espacios_curriculares = [col for col in df.columns if col.startswith("ESPACIO CURRICULAR")]
        valores_espacios_curriculares = df[columnas_espacios_curriculares]
        # Verificar si la materia está en los valores de espacios curriculares
        materia_en_espacios = materia in valores_espacios_curriculares.values
        columnas_cursos = [col for col in df.columns if col.startswith("CURSO")]
        valores_cursos = df[columnas_cursos]
        #Verificar si el curso está en los valores de cursos
        curso_en_curso = curso in valores_cursos.values
        # Verificar si el alumno ya está registrado en la misma materia
        if df[(df["DNI"] == int(dni)) & materia_en_espacios & curso_en_curso].shape[0] > 0:
            messagebox.showerror("Error", "Este alumno ya está registrado en esta materia.")
            return

        # Encontrar el índice del alumno
        index = df[df["DNI"] == int(dni)].index[0]

        # Buscar el número de la siguiente columna disponible para cada tipo de dato
        columna_condicion = next(
            (col for col in df.columns if col.startswith("CONDICION") and pd.isna(df.at[index, col])), None)
        columna_curso = next(
            (col for col in df.columns if col.startswith("CURSO") and pd.isna(df.at[index, col])), None)
        columna_materia = next(
            (col for col in df.columns if col.startswith("ESPACIO CURRICULAR") and pd.isna(df.at[index, col])), None)

        # Si no existen columnas disponibles, crear nuevas
        if columna_condicion is None:
            columna_condicion = f"CONDICION {len([col for col in df.columns if col.startswith('CONDICION')]) + 1}"
            df[columna_condicion] = pd.NA
        if columna_curso is None:
            columna_curso = f"CURSO {len([col for col in df.columns if col.startswith('CURSO')]) + 1}"
            df[columna_curso] = pd.NA
        if columna_materia is None:
            columna_materia = f"ESPACIO CURRICULAR {len([col for col in df.columns if col.startswith('ESPACIO CURRICULAR')]) + 1}"
            df[columna_materia] = pd.NA

        # Asignar los nuevos valores en las columnas correspondientes
        df.at[index, columna_condicion] = condicion
        df.at[index, columna_curso] = curso
        df.at[index, columna_materia] = materia

    else:
        # Crear un nuevo registro con columnas iniciales
        nuevo_registro = {
            "N°": str(len(df) + 1),
            "ALUMNO": nombre,
            "DNI": dni,
            "MODALIDAD": especialidad,
            "ESPACIO CURRICULAR": materia,
            "CURSO": curso,
            "CONDICION": condicion
        }

        # Completar las columnas dinámicas con valores vacíos
        for col in df.columns:
            if col not in nuevo_registro:
                nuevo_registro[col] = pd.NA

        # Añadir el nuevo registro al DataFrame
        df = pd.concat([df, pd.DataFrame([nuevo_registro])], ignore_index=True)

    # Intentar guardar el DataFrame en el archivo Excel
    try:
        # Sobreescribir el archivo Excel con openpyxl
        with pd.ExcelWriter(archivo_excel, engine="openpyxl", mode="w") as writer:
            df.to_excel(writer, index=False)
        generar_actas()
        ajustar_columnas_automatically(archivo_excel, "PERMISOS")
        ajustar_columnas_automatically(archivo_excel, "MMO")
        ajustar_columnas_automatically(archivo_excel, "TEM")
        messagebox.showinfo("Éxito",
                             "Alumno registrado correctamente.")
        limpiar_campos()
        cargar_datos()
    except PermissionError:
        messagebox.showerror("Error",
                             "El archivo Excel está abierto. Por favor, cierre el archivo y luego registre al alumno.")

def buscar_alumno_nombre(event=None):
    # Búsqueda parcial
    nombre_buscado = entrada_nombre.get().strip()

    # Leer el archivo excel
    df = pd.read_excel(archivo_excel)

    for item in tabla.get_children():
        tabla.delete(item)
    if nombre_buscado == "":
        # Si el campo está vacío, mostrar todos los registros
        for _, row in df.iterrows():
            tabla.insert("", tk.END, values=list(row.split(" (")[0]))
        return
    # Buscar coincidencias parciales en el campo "ALUMNO"
    df_filtrado = df[df["ALUMNO"].astype(str).str.contains(nombre_buscado, case=False, na=False)]

    if not df_filtrado.empty:
        # Mostrar cada registro del alumno encontrado en la tabla
        for _, row in df_filtrado.iterrows():
            tabla.insert("", tk.END, values=list(row))
    else:
        # Si no se encuentra ninguna coincidencia
        messagebox.showerror("Error", "No se encontró ningún alumno con el nombre ingresado.")

def buscar_alumno(event=None):
    dni_buscado = entrada_dni.get().strip()

    # Validar si el DNI ingresado es un número o está vacío
    if not dni_buscado.isdigit() and dni_buscado != "":
        messagebox.showerror("Error", "Por favor, ingrese un DNI válido (solo números).")
        return

    # Leer el archivo Excel
    df = pd.read_excel(archivo_excel)

    # Limpiar la tabla antes de insertar nuevos datos
    for item in tabla.get_children():
        tabla.delete(item)

    if dni_buscado == "":
        # Si el campo está vacío, mostrar todos los registros
        for _, row in df.iterrows():
            tabla.insert("", tk.END, values=list(row))
        return

    # Buscar coincidencias parciales en el campo "DNI"
    df_filtrado = df[df["DNI"].astype(str).str.contains(dni_buscado, case=False, na=False)]

    if not df_filtrado.empty:
        # Mostrar cada registro del alumno encontrado en la tabla
        for _, row in df_filtrado.iterrows():
            tabla.insert("", tk.END, values=list(row))
    else:
        # Si no se encuentra ninguna coincidencia
        messagebox.showerror("Error", "No se encontró ningún alumno con el DNI ingresado.")

def actualizar_materias(event):
    modalidad = combobox_especialidad.get().strip().upper()
    if modalidad == "MMO":
        combobox_materia['values'] = materias_mmo
    elif modalidad == "TEM":
        combobox_materia['values'] = materias_tem
    else:
        combobox_materia['values'] = []  # Si no se selecciona ninguna modalidad

def buscar_materias(event):
    # Obtener el texto ingresado por el usuario
    texto = combobox_materia.get().upper()  # Convertir a mayúsculas para que sea insensible a mayúsculas/minúsculas

    # Filtrar las materias según el texto ingresado
    if combobox_especialidad.get() == "MMO":
        lista_materias = [materia for materia in materias_mmo if texto in materia.upper()]
    elif combobox_especialidad.get() == "TEM":
        lista_materias = [materia for materia in materias_tem if texto in materia.upper()]
    else:
        lista_materias = []

    # Actualizar el ComboBox con las materias filtradas
    combobox_materia['values'] = lista_materias

    # Si hay coincidencias, no seleccionamos nada automáticamente
    if texto == "":
        combobox_materia.set("")  # Si el campo de búsqueda está vacío, dejamos el combobox vacío


# Función para limpiar los campos del formulario
def limpiar_campos():
    combobox_curso.set("")  # Restablecer combobox
    combobox_materia.delete(0, tk.END)
    combobox_condicion.set("")  # Restablecer combobox

# Validación para que solo se permitan números en el campo de DNI
def validar_dni(caracter):
    return caracter.isdigit()

# Función para cargar datos desde un archivo Excel y actualizar la tabla
def cargar_datos():
    archivo_excel = "acta_de_examen.xlsx"
    try:
        # Leer archivo Excel
        df = pd.read_excel(archivo_excel)

        # Limpiar la tabla existente
        for item in tabla.get_children():
            tabla.delete(item)
        tabla["columns"] = list(df.columns)  # Obtener nombres de columnas

        # Configurar encabezados de columnas en el Treeview
        for col in df.columns:
            tabla.heading(col, text=col)
            tabla.column(col, anchor="center", width=150)  # Ajustar ancho de columna

        # Agregar filas a la tabla
        for _, row in df.iterrows():
            tabla.insert("", "end", values=list(row))
    except FileNotFoundError:
        pass

# Configurar la ventana principal
ventana = tk.Tk()
ventana.title("Registro de Acta de Examen")
ventana.geometry("1080x720")  # Aumentar tamaño para incluir la tabla
ventana.configure(bg="lightblue")

# Crear un solo frame para la tabla
frame_tabla = tk.Frame(ventana)
frame_tabla.grid(row=8, column=0, columnspan=4, padx=10, pady=10, sticky="nsew")  # Ocupa todas las columnas

# Crear estilo para el Treeview
style = ttk.Style()
style.configure("Treeview", font=("Calibri", 11))  # Tamaño de las celdas
style.configure("Treeview.Heading", font=("Calibri", 11, "bold"))  # Tamaño de los encabezados

# Crear la tabla (Treeview)
tabla = ttk.Treeview(
    frame_tabla,
    columns=columnas,
    show="headings",
    height=10
)

# Configurar las barras de desplazamiento
scroll_y = ttk.Scrollbar(frame_tabla, orient="vertical", command=tabla.yview)
scroll_x = ttk.Scrollbar(frame_tabla, orient="horizontal", command=tabla.xview)

tabla.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)

# Empaquetar la tabla
tabla.grid(row=0, column=0, sticky="nsew")  # Usar grid en lugar de pack para una mejor distribución

# Configurar las barras de desplazamiento
scroll_y.grid(row=0, column=1, sticky="ns")
scroll_x.grid(row=1, column=0, sticky="ew")

ventana.columnconfigure(0, weight=0)  # Opcional si no se necesita expandir
ventana.columnconfigure(1, weight=1)  # Columna del frame de la tabla
ventana.columnconfigure(2, weight=1)
ventana.columnconfigure(3, weight=1)  # Todas las columnas contribuyen al ancho

ventana.rowconfigure(8, weight=1)  # Permitir que la fila de la tabla se expanda
frame_tabla.grid_rowconfigure(0, weight=1)  # Hacer que la tabla se expanda dentro del frame
frame_tabla.grid_columnconfigure(0, weight=1)

# Esto asegura que la tabla ocupe el espacio necesario para permitir el desplazamiento
cargar_datos()
# Configurar encabezados de la tabla
for col in columnas:
    tabla.heading(col, text=col)
    tabla.column(col, width=150, anchor="center")  # Ajustar ancho de las columnas

# Configuración para expansión del frame_tabla
ventana.columnconfigure(0, weight=1)
ventana.rowconfigure(8, weight=1)

# Generar lista de cursos (sin el símbolo ° repetido)
cursos = [f"{año}°{div}°" for año in range(1, 8) for div in range(1, 3)]

# Etiquetas y campos de entrada con alineación centrada
tk.Label(ventana, text="Nombre del Alumno:", anchor="e", font=("Calibri", 11)).grid(row=0, column=0, padx=10, pady=10, sticky="e")
entrada_nombre = tk.Entry(ventana, justify="center", font=("Calibri", 11))
entrada_nombre.grid(row=0, column=1, padx=10, pady=10, sticky="ew")

tk.Label(ventana, text="DNI:", anchor="e", font=("Calibri", 11)).grid(row=1, column=0, padx=10, pady=10, sticky="e")
vcmd = (ventana.register(validar_dni), "%S")  # Validación de entrada
entrada_dni = tk.Entry(ventana, validate="key", validatecommand=vcmd, justify="center", font=("Calibri", 11))
# Evento para buscar al alumno cuando el usuario escribe
entrada_dni.grid(row=1, column=1, padx=10, pady=10, sticky="ew")

tk.Label(ventana, text="Curso:", anchor="e", font=("Calibri", 11)).grid(row=4, column=0, padx=10, pady=10, sticky="e")
combobox_curso = ttk.Combobox(ventana, values=cursos, state="readonly", justify="center", font=("Calibri", 11))
combobox_curso.grid(row=4, column=1, padx=10, pady=10, sticky="ew")
combobox_curso.set("Seleccione un curso...")

tk.Label(ventana, text="Modalidad:", anchor="e", font=("Calibri", 11)).grid(row=3, column=0, padx=10, pady=10, sticky="e")
combobox_especialidad = ttk.Combobox(ventana, values=["MMO", "TEM"], state="readonly", justify="center", font=("Calibri", 11))
combobox_especialidad.bind("<<ComboboxSelected>>", actualizar_materias)
combobox_especialidad.grid(row=3, column=1, padx=10, pady=10, sticky="ew")
combobox_especialidad.set("Seleccione una especialidad...")

tk.Label(ventana, text="Materia:", anchor="e", font=("Calibri", 11)).grid(row=5, column=0, padx=10, pady=10, sticky="e")
combobox_materia = ttk.Combobox(ventana, justify="center", font=("Calibri", 11))
combobox_materia.bind("<KeyRelease>", buscar_materias)  # Buscar materias mientras se escribe
combobox_materia.grid(row=5, column=1, padx=10, pady=10, sticky="ew")
combobox_materia.set("Seleccione una materia...")

tk.Label(ventana, text="Condición:", anchor="e", font=("Calibri", 11)).grid(row=6, column=0, padx=10, pady=10, sticky="e")
combobox_condicion = ttk.Combobox(ventana, values=["LIBRE", "REGULAR"], state="readonly", justify="center", font=("Calibri", 11))
combobox_condicion.grid(row=6, column=1, padx=10, pady=10, sticky="ew")
combobox_condicion.set("Seleccione una condicion...")

# Botones
from PIL import Image, ImageTk  # Pillow
imagen_original_lupa = Image.open("recursos\image\loupe.png")
imagen_redimensionada_lupa = imagen_original_lupa.resize((30, 30))
lupa = ImageTk.PhotoImage(imagen_redimensionada_lupa)
imagen_original_word = Image.open("recursos/image/google-docs.png")
imagen_redimensionada_word = imagen_original_word.resize((30, 30))
word = ImageTk.PhotoImage(imagen_redimensionada_word)
imagen_original_excel = Image.open("recursos/image/xls.png")
imagen_redimensionada_excel = imagen_original_excel.resize((30, 30))
excel = ImageTk.PhotoImage(imagen_redimensionada_excel)
boton_registrar = tk.Button(ventana, text="Registrar", command=registrar, font=("Calibri", 11, "bold"), bg="white", image=excel, compound="left",)
boton_registrar.grid(row=6, column=2, padx=10, pady=10, sticky="ew")
boton_generar_actas = tk.Button(ventana, text="Generar Actas", command=generar_actas_imprimir, font=("Calibri", 11, "bold"), bg="white", image=word, compound="left",)
boton_generar_actas.grid(row=7, column=2, padx=10, pady=10, sticky="ew")
boton_generar_permisos = tk.Button(ventana, text="Generar Permisos", command=generar_permisos, font=("Calibri", 11, "bold"), bg="white", image=word, compound="left",)
boton_generar_permisos.grid(row=7, column=3, padx=10, pady=10, sticky="ew")
boton_buscar_dni = tk.Button(ventana, text="Buscar por DNI", command=buscar_alumno, font = ("Calibri", 11, "bold"), bg="white", image=lupa, compound="left")
boton_buscar_dni.grid(row=1, column=2, padx=10, pady=10, sticky="ew")
boton_buscar_alumno = tk.Button(ventana, text="Buscar por Nombre", command=buscar_alumno_nombre, font = ("Calibri", 11, "bold"), bg="white", image=lupa, compound="left")
boton_buscar_alumno.grid(row=0, column=2, padx=10, pady=10, sticky="ew")
# Cargar el archivo Excel
wb = load_workbook(archivo_excel)
ws = wb.active

# Desbloquear la hoja para ediciones
ws.protection = SheetProtection(sheet=False)
ajustar_columnas_automatically(archivo_excel, "PERMISOS")
ajustar_columnas_automatically(archivo_excel, "MMO")
ajustar_columnas_automatically(archivo_excel, "TEM")

# Reproteger la hoja después de editar
wb = load_workbook(archivo_excel)
ws = wb.active
ws.protection = SheetProtection(sheet=True, password=contraseña)

# Guardar el archivo nuevamente con protección
wb.save(archivo_excel)

# Configurar las columnas de la cuadrícula para que sean expandibles
ventana.columnconfigure(0, weight=1)  # Columna de etiquetas (ancla derecha)
ventana.columnconfigure(1, weight=3)  # Columna de campos de entrada (expansible)
ventana.columnconfigure(2, weight=1)  # Columna de botones (opcional)

# Iniciar el bucle principal de la interfaz gráfica
ventana.mainloop()
